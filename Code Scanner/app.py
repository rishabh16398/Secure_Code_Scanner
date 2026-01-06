from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, Response, stream_with_context
from werkzeug.utils import secure_filename
from pathlib import Path
import tempfile, shutil, uuid
import openpyxl
import json
import time
from queue import Queue
from threading import Thread

from scanners import run_all_scanners, build_sast_summary, build_dep_summary
from docx import Document  # needs python-docx
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment

app = Flask(__name__)
app.secret_key = "change-me"

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
REPORT_DIR = BASE_DIR / "reports"
UPLOAD_DIR.mkdir(exist_ok=True)
REPORT_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {"zip"}

# Changed to store multiple runs
ALL_RUNS = []  # List of dicts, each containing results for one run

# Progress tracking
progress_queues = {}  # scan_id -> Queue for progress updates


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def generate_docx_report(project_name, sast_summary, dep_summary, out_path: Path):
    doc = Document()

    doc.add_heading("Security Scan Report", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph()

    doc.add_heading("SAST CWE Summary", level=2)
    if sast_summary:
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "CWE"
        hdr[1].text = "Severity"
        hdr[2].text = "Occurrences"
        hdr[3].text = "Found by"
        hdr[4].text = "Example files / lines"

        for row in sast_summary:
            cells = table.add_row().cells
            cells[0].text = row.get("cwe", "")
            cells[1].text = row.get("severity", "")
            cells[2].text = str(row.get("occurrences", 0))
            cells[3].text = ", ".join(row.get("scanners", []))

            examples = row.get("examples", [])
            if examples:
                lines = []
                for ex in examples:
                    file_ = ex.get("file")
                    line = ex.get("line")
                    scanner = ex.get("scanner")
                    lines.append(f"{file_}:{line} ({scanner})")
                cells[4].text = "\n".join(lines)
            else:
                cells[4].text = "-"

    else:
        doc.add_paragraph("No SAST findings.")

    doc.add_page_break()

    doc.add_heading("Dependency Vulnerability Summary", level=2)
    if dep_summary:
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "CVE / ID"
        hdr[1].text = "Severity"
        hdr[2].text = "Occurrences"
        hdr[3].text = "Found by"
        hdr[4].text = "Packages"

        for row in dep_summary:
            cells = table.add_row().cells
            cells[0].text = row.get("cve", "")
            cells[1].text = row.get("severity", "")
            cells[2].text = str(row.get("occurrences", 0))
            cells[3].text = ", ".join(row.get("scanners", []))
            cells[4].text = ", ".join(row.get("packages", []))
    else:
        doc.add_paragraph("No dependency findings.")

    doc.save(out_path)


def generate_comparison_excel(all_runs_data, out_path: Path):
    """
    Generate an Excel file comparing CWEs across multiple runs.
    
    Format matches user's manual Excel exactly:
    - Column A: CWE ID
    - Column B: CWE Name
    - Columns C onwards: Run 1, Run 2, ... Run N
    - Then: Runs Found, Total Runs, All Tools Used, Tool Details by Run
    - Cells: ✓ if CWE found in that run, empty otherwise
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "CWE Comparison"
    
    # CWE name mapping (comprehensive list)
    cwe_names = {
        "CWE-22": "Improper Limitation of a Pathname to a Restricted Directory ('Path Traversal')",
        "CWE-78": "Improper Neutralization of Special Elements used in an OS Command ('OS Command Injection')",
        "CWE-79": "Improper Neutralization of Input During Web Page Generation ('Cross-site Scripting')",
        "CWE-89": "Improper Neutralization of Special Elements used in an SQL Command ('SQL Injection')",
        "CWE-94": "Improper Control of Generation of Code ('Code Injection')",
        "CWE-119": "Improper Restriction of Operations within the Bounds of a Memory Buffer",
        "CWE-125": "Out-of-bounds Read",
        "CWE-190": "Integer Overflow or Wraparound",
        "CWE-200": "Exposure of Sensitive Information to an Unauthorized Actor",
        "CWE-250": "Execution with Unnecessary Privileges",
        "CWE-287": "Improper Authentication",
        "CWE-295": "Improper Certificate Validation",
        "CWE-297": "Improper Validation of Certificate with Host Mismatch",
        "CWE-306": "Missing Authentication for Critical Function",
        "CWE-312": "Cleartext Storage of Sensitive Information",
        "CWE-319": "Cleartext Transmission of Sensitive Information",
        "CWE-327": "Use of a Broken or Risky Cryptographic Algorithm",
        "CWE-338": "Use of Cryptographically Weak Pseudo-Random Number Generator (PRNG)",
        "CWE-377": "Insecure Temporary File",
        "CWE-400": "Uncontrolled Resource Consumption",
        "CWE-502": "Deserialization of Untrusted Data",
        "CWE-521": "Weak Password Requirements",
        "CWE-601": "URL Redirection to Untrusted Site ('Open Redirect')",
        "CWE-611": "Improper Restriction of XML External Entity Reference",
        "CWE-614": "Sensitive Cookie in HTTPS Session Without 'Secure' Attribute",
        "CWE-732": "Incorrect Permission Assignment for Critical Resource",
        "CWE-776": "Unrestricted Recursion",
        "CWE-798": "Use of Hard-coded Credentials",
        "CWE-917": "Improper Neutralization of Special Elements used in an Expression Language Statement ('Expression Language Injection')",
    }
    
    # Collect all unique CWEs across all runs with detailed info
    cwe_data = {}  # cwe -> {runs: set(), tools_by_run: {run_idx: set()}}
    
    for run_idx, run_data in enumerate(all_runs_data):
        sast_summary = run_data.get("sast_summary", [])
        for item in sast_summary:
            cwe = item.get("cwe")
            if cwe and cwe != "CWE-UNKNOWN":
                if cwe not in cwe_data:
                    cwe_data[cwe] = {
                        "runs": set(),
                        "tools_by_run": {}
                    }
                
                cwe_data[cwe]["runs"].add(run_idx)
                
                # Track which tools found this CWE in this run
                if run_idx not in cwe_data[cwe]["tools_by_run"]:
                    cwe_data[cwe]["tools_by_run"][run_idx] = set()
                
                # Get tools that found this CWE
                scanners = item.get("scanners", [])
                for scanner in scanners:
                    cwe_data[cwe]["tools_by_run"][run_idx].add(scanner)
    
    # Sort CWEs
    sorted_cwes = sorted(cwe_data.keys())
    
    # Define styles
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    check_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    center_alignment = Alignment(horizontal="center", vertical="center")
    wrap_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    
    num_runs = len(all_runs_data)
    
    # Write headers
    headers = ["CWE ID", "CWE Name"]
    for idx in range(1, num_runs + 1):
        headers.append(f"Run {idx}")
    headers.extend(["Runs Found", "Total Runs", "All Tools Used", "Tool Details by Run"])
    
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_alignment
    
    # Write CWE data rows
    for row_idx, cwe in enumerate(sorted_cwes, start=2):
        data = cwe_data[cwe]
        
        # Column A: CWE ID
        ws.cell(row=row_idx, column=1, value=cwe)
        
        # Column B: CWE Name
        cwe_name = cwe_names.get(cwe, "")
        ws.cell(row=row_idx, column=2, value=cwe_name)
        
        # Columns for each run (checkmarks)
        for run_idx in range(num_runs):
            col = 3 + run_idx
            cell = ws.cell(row=row_idx, column=col)
            if run_idx in data["runs"]:
                cell.value = "✓"
                cell.fill = check_fill
                cell.alignment = center_alignment
                cell.font = Font(bold=True, size=14)
        
        # Runs Found column
        runs_found_list = sorted([f"Run {i+1}" for i in data["runs"]])
        runs_found_str = ", ".join(runs_found_list)
        ws.cell(row=row_idx, column=3 + num_runs, value=runs_found_str)
        
        # Total Runs column
        total_runs = len(data["runs"])
        ws.cell(row=row_idx, column=4 + num_runs, value=total_runs).alignment = center_alignment
        
        # All Tools Used column
        all_tools = set()
        for tools in data["tools_by_run"].values():
            all_tools.update(tools)
        all_tools_str = ", ".join(sorted(all_tools))
        ws.cell(row=row_idx, column=5 + num_runs, value=all_tools_str)
        
        # Tool Details by Run column
        tool_details = []
        for run_idx in sorted(data["tools_by_run"].keys()):
            tools = sorted(data["tools_by_run"][run_idx])
            tool_details.append(f"Run {run_idx + 1}: {', '.join(tools)}")
        tool_details_str = " | ".join(tool_details)
        cell = ws.cell(row=row_idx, column=6 + num_runs, value=tool_details_str)
        cell.alignment = wrap_alignment
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 12  # CWE ID
    ws.column_dimensions['B'].width = 60  # CWE Name
    
    for col_idx in range(3, 3 + num_runs):
        col_letter = openpyxl.utils.get_column_letter(col_idx)
        ws.column_dimensions[col_letter].width = 10  # Run columns
    
    ws.column_dimensions[openpyxl.utils.get_column_letter(3 + num_runs)].width = 30  # Runs Found
    ws.column_dimensions[openpyxl.utils.get_column_letter(4 + num_runs)].width = 12  # Total Runs
    ws.column_dimensions[openpyxl.utils.get_column_letter(5 + num_runs)].width = 25  # All Tools Used
    ws.column_dimensions[openpyxl.utils.get_column_letter(6 + num_runs)].width = 80  # Tool Details by Run
    
    wb.save(out_path)


def generate_detailed_count_excel(all_runs_data, out_path: Path):
    """
    NEW FEATURE: Generate Excel with FILE COUNTS instead of checkmarks
    
    Instead of ✓, shows the NUMBER OF FILES that CWE was found in
    Example: CWE-79 in Run 1 found in 20 files -> shows "20"
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "CWE Detailed Counts"
    
    # COMPREHENSIVE CWE name mapping (100+ CWEs)
    cwe_names = {
        # Input Validation
        "CWE-20": "Improper Input Validation",
        "CWE-74": "Improper Neutralization of Special Elements in Output Used by a Downstream Component ('Injection')",
        "CWE-75": "Failure to Sanitize Special Elements into a Different Plane (Special Element Injection)",
        "CWE-79": "Improper Neutralization of Input During Web Page Generation ('Cross-site Scripting')",
        "CWE-89": "Improper Neutralization of Special Elements used in an SQL Command ('SQL Injection')",
        "CWE-91": "XML Injection (aka Blind XPath Injection)",
        "CWE-94": "Improper Control of Generation of Code ('Code Injection')",
        "CWE-95": "Improper Neutralization of Directives in Dynamically Evaluated Code ('Eval Injection')",
        "CWE-96": "Improper Neutralization of Directives in Statically Saved Code ('Static Code Injection')",
        "CWE-97": "Improper Neutralization of Server-Side Includes (SSI) Within a Web Page",
        "CWE-98": "Improper Control of Filename for Include/Require Statement in PHP Program ('PHP Remote File Inclusion')",
        
        # Path Traversal & File Operations
        "CWE-22": "Improper Limitation of a Pathname to a Restricted Directory ('Path Traversal')",
        "CWE-23": "Relative Path Traversal",
        "CWE-36": "Absolute Path Traversal",
        "CWE-73": "External Control of File Name or Path",
        "CWE-434": "Unrestricted Upload of File with Dangerous Type",
        "CWE-59": "Improper Link Resolution Before File Access ('Link Following')",
        "CWE-377": "Insecure Temporary File",
        
        # Command Injection
        "CWE-78": "Improper Neutralization of Special Elements used in an OS Command ('OS Command Injection')",
        "CWE-77": "Improper Neutralization of Special Elements used in a Command ('Command Injection')",
        
        # Authentication & Access Control
        "CWE-287": "Improper Authentication",
        "CWE-306": "Missing Authentication for Critical Function",
        "CWE-285": "Improper Authorization",
        "CWE-284": "Improper Access Control",
        "CWE-862": "Missing Authorization",
        "CWE-863": "Incorrect Authorization",
        "CWE-276": "Incorrect Default Permissions",
        "CWE-732": "Incorrect Permission Assignment for Critical Resource",
        "CWE-250": "Execution with Unnecessary Privileges",
        "CWE-269": "Improper Privilege Management",
        
        # Cryptography
        "CWE-327": "Use of a Broken or Risky Cryptographic Algorithm",
        "CWE-328": "Use of Weak Hash",
        "CWE-326": "Inadequate Encryption Strength",
        "CWE-321": "Use of Hard-coded Cryptographic Key",
        "CWE-322": "Key Exchange without Entity Authentication",
        "CWE-323": "Reusing a Nonce, Key Pair in Encryption",
        "CWE-324": "Use of a Key Past its Expiration Date",
        "CWE-325": "Missing Cryptographic Step",
        "CWE-329": "Generation of Predictable IV with CBC Mode",
        "CWE-338": "Use of Cryptographically Weak Pseudo-Random Number Generator (PRNG)",
        "CWE-330": "Use of Insufficiently Random Values",
        "CWE-331": "Insufficient Entropy",
        "CWE-335": "Incorrect Usage of Seeds in Pseudo-Random Number Generator (PRNG)",
        "CWE-336": "Same Seed in Pseudo-Random Number Generator (PRNG)",
        "CWE-337": "Predictable Seed in Pseudo-Random Number Generator (PRNG)",
        
        # SSL/TLS
        "CWE-295": "Improper Certificate Validation",
        "CWE-296": "Improper Following of a Certificate's Chain of Trust",
        "CWE-297": "Improper Validation of Certificate with Host Mismatch",
        "CWE-298": "Improper Validation of Certificate Expiration",
        "CWE-299": "Improper Check for Certificate Revocation",
        
        # Sensitive Data
        "CWE-200": "Exposure of Sensitive Information to an Unauthorized Actor",
        "CWE-209": "Generation of Error Message Containing Sensitive Information",
        "CWE-215": "Insertion of Sensitive Information Into Debugging Code",
        "CWE-312": "Cleartext Storage of Sensitive Information",
        "CWE-313": "Cleartext Storage in a File or on Disk",
        "CWE-314": "Cleartext Storage in the Registry",
        "CWE-315": "Cleartext Storage of Sensitive Information in a Cookie",
        "CWE-316": "Cleartext Storage of Sensitive Information in Memory",
        "CWE-317": "Cleartext Storage of Sensitive Information in GUI",
        "CWE-318": "Cleartext Storage of Sensitive Information in Executable",
        "CWE-319": "Cleartext Transmission of Sensitive Information",
        "CWE-321": "Use of Hard-coded Cryptographic Key",
        "CWE-798": "Use of Hard-coded Credentials",
        "CWE-259": "Use of Hard-coded Password",
        "CWE-257": "Storing Passwords in a Recoverable Format",
        
        # Session Management
        "CWE-384": "Session Fixation",
        "CWE-613": "Insufficient Session Expiration",
        "CWE-614": "Sensitive Cookie in HTTPS Session Without 'Secure' Attribute",
        "CWE-1004": "Sensitive Cookie Without 'HttpOnly' Flag",
        "CWE-565": "Reliance on Cookies without Validation and Integrity Checking",
        
        # CSRF & Redirects
        "CWE-352": "Cross-Site Request Forgery (CSRF)",
        "CWE-601": "URL Redirection to Untrusted Site ('Open Redirect')",
        
        # XML & XXE
        "CWE-611": "Improper Restriction of XML External Entity Reference",
        "CWE-827": "Improper Control of Document Type Definition",
        
        # Deserialization
        "CWE-502": "Deserialization of Untrusted Data",
        "CWE-915": "Improperly Controlled Modification of Dynamically-Determined Object Attributes",
        
        # Resource Management
        "CWE-400": "Uncontrolled Resource Consumption",
        "CWE-404": "Improper Resource Shutdown or Release",
        "CWE-770": "Allocation of Resources Without Limits or Throttling",
        "CWE-771": "Missing Reference to Active Allocated Resource",
        "CWE-772": "Missing Release of Resource after Effective Lifetime",
        "CWE-775": "Missing Release of File Descriptor or Handle after Effective Lifetime",
        "CWE-776": "Unrestricted Recursion",
        "CWE-834": "Excessive Iteration",
        
        # Memory Safety
        "CWE-119": "Improper Restriction of Operations within the Bounds of a Memory Buffer",
        "CWE-120": "Buffer Copy without Checking Size of Input ('Classic Buffer Overflow')",
        "CWE-121": "Stack-based Buffer Overflow",
        "CWE-122": "Heap-based Buffer Overflow",
        "CWE-125": "Out-of-bounds Read",
        "CWE-787": "Out-of-bounds Write",
        "CWE-416": "Use After Free",
        "CWE-415": "Double Free",
        "CWE-476": "NULL Pointer Dereference",
        "CWE-401": "Missing Release of Memory after Effective Lifetime",
        "CWE-911": "Improper Update of Reference Count",
        
        # Integer & Numeric Errors
        "CWE-190": "Integer Overflow or Wraparound",
        "CWE-191": "Integer Underflow (Wrap or Wraparound)",
        "CWE-680": "Integer Overflow to Buffer Overflow",
        "CWE-681": "Incorrect Conversion between Numeric Types",
        "CWE-682": "Incorrect Calculation",
        "CWE-369": "Divide By Zero",
        
        # Race Conditions
        "CWE-362": "Concurrent Execution using Shared Resource with Improper Synchronization ('Race Condition')",
        "CWE-367": "Time-of-check Time-of-use (TOCTOU) Race Condition",
        "CWE-364": "Signal Handler Race Condition",
        
        # Logic Errors
        "CWE-670": "Always-Incorrect Control Flow Implementation",
        "CWE-571": "Expression is Always True",
        "CWE-570": "Expression is Always False",
        "CWE-561": "Dead Code",
        "CWE-489": "Active Debug Code",
        "CWE-501": "Trust Boundary Violation",
        
        # Logging & Monitoring
        "CWE-117": "Improper Output Neutralization for Logs",
        "CWE-532": "Insertion of Sensitive Information into Log File",
        "CWE-533": "DEPRECATED: Information Exposure Through Server Log Files",
        
        # Configuration
        "CWE-1188": "Insecure Default Initialization of Resource",
        "CWE-426": "Untrusted Search Path",
        "CWE-427": "Uncontrolled Search Path Element",
        "CWE-829": "Inclusion of Functionality from Untrusted Control Sphere",
        "CWE-830": "Inclusion of Web Functionality from an Untrusted Source",
        
        # Code Quality
        "CWE-477": "Use of Obsolete Function",
        "CWE-478": "Missing Default Case in Multiple Condition Expression",
        "CWE-479": "Signal Handler Use of a Non-reentrant Function",
        "CWE-480": "Use of Incorrect Operator",
        "CWE-483": "Incorrect Block Delimitation",
        "CWE-484": "Omitted Break Statement in Switch",
        
        # Expression Language Injection
        "CWE-917": "Improper Neutralization of Special Elements used in an Expression Language Statement ('Expression Language Injection')",
        
        # Password & Auth
        "CWE-521": "Weak Password Requirements",
        "CWE-916": "Use of Password Hash With Insufficient Computational Effort",
        
        # Regex
        "CWE-1333": "Inefficient Regular Expression Complexity",
        
        # Prototype Pollution
        "CWE-1321": "Improperly Controlled Modification of Object Prototype Attributes ('Prototype Pollution')",
        
        # Server-Side Request Forgery
        "CWE-918": "Server-Side Request Forgery (SSRF)",
        
        # Null Byte Injection
        "CWE-158": "Improper Neutralization of Null Byte or NUL Character",
        
        # Format String
        "CWE-134": "Use of Externally-Controlled Format String",
        
        # Uncontrolled Format String
        "CWE-134": "Use of Externally-Controlled Format String",
        
        # Information Disclosure
        "CWE-203": "Observable Discrepancy",
        "CWE-208": "Observable Timing Discrepancy",
        
        # Missing Support
        "CWE-353": "Missing Support for Integrity Check",
        
        # ZIP vulnerabilities
        "CWE-409": "Improper Handling of Highly Compressed Data (Data Amplification)",
        "CWE-410": "Insufficient Resource Pool",
    }
    
    # Collect CWE data with FILE COUNTS per run
    # FIXED: Use RAW scanner results instead of summary (which only has 5 examples)
    cwe_data = {}  # cwe -> {run_idx: {total_count, unique_files, tools}}
    
    for run_idx, run_data in enumerate(all_runs_data):
        # Get RAW SAST results (not summary which only has 5 examples!)
        sast_results = run_data.get("results", {}).get("sast", {})
        
        # Process each scanner's results
        for scanner_name, findings in sast_results.items():
            for finding in findings:
                cwe = finding.get("cwe")
                if cwe and cwe != "CWE-UNKNOWN":
                    if cwe not in cwe_data:
                        cwe_data[cwe] = {}
                    
                    if run_idx not in cwe_data[cwe]:
                        cwe_data[cwe][run_idx] = {
                            "total_count": 0,  # Total occurrences
                            "unique_files": set(),  # Unique files
                            "tools": set()
                        }
                    
                    # Count this occurrence
                    cwe_data[cwe][run_idx]["total_count"] += 1
                    
                    # Add unique file
                    file_path = finding.get("file")
                    if file_path:
                        cwe_data[cwe][run_idx]["unique_files"].add(file_path)
                    
                    # Track tool
                    cwe_data[cwe][run_idx]["tools"].add(scanner_name)
    
    # Sort CWEs
    sorted_cwes = sorted(cwe_data.keys())
    
    # Define styles
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    count_fill = PatternFill(start_color="FFE699", end_color="FFE699", fill_type="solid")  # Yellow for counts
    center_alignment = Alignment(horizontal="center", vertical="center")
    wrap_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    
    num_runs = len(all_runs_data)
    
    # Write headers - TWO columns per run (Total and Unique)
    headers_row1 = ["CWE ID", "CWE Name"]
    
    # Add run headers (each run gets 2 columns)
    for idx in range(1, num_runs + 1):
        headers_row1.append(f"Run {idx} Total")
        headers_row1.append(f"Run {idx} Unique")
    
    headers_row1.extend(["Grand Total", "Grand Unique", "All Tools Used"])
    
    for col_idx, header in enumerate(headers_row1, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_alignment
    
    # Write CWE data rows
    for row_idx, cwe in enumerate(sorted_cwes, start=2):
        data = cwe_data[cwe]
        
        # Column A: CWE ID
        ws.cell(row=row_idx, column=1, value=cwe)
        
        # Column B: CWE Name
        cwe_name = cwe_names.get(cwe, "")
        ws.cell(row=row_idx, column=2, value=cwe_name)
        
        grand_total_count = 0
        grand_unique_files = set()
        
        # Columns for each run (TWO columns: Total and Unique)
        col = 3
        for run_idx_col in range(num_runs):
            if run_idx_col in data:
                # Total Count
                total_count = data[run_idx_col]["total_count"]
                cell = ws.cell(row=row_idx, column=col)
                cell.value = total_count
                cell.fill = PatternFill(start_color="FFE699", end_color="FFE699", fill_type="solid")  # Yellow
                cell.alignment = center_alignment
                cell.font = Font(bold=True, size=11)
                
                # Unique File Count
                unique_count = len(data[run_idx_col]["unique_files"])
                cell = ws.cell(row=row_idx, column=col+1)
                cell.value = unique_count
                cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")  # Green
                cell.alignment = center_alignment
                cell.font = Font(bold=True, size=11)
                
                grand_total_count += total_count
                grand_unique_files.update(data[run_idx_col]["unique_files"])
            else:
                # No data for this run
                ws.cell(row=row_idx, column=col, value=0).alignment = center_alignment
                ws.cell(row=row_idx, column=col+1, value=0).alignment = center_alignment
            
            col += 2
        
        # Grand Total column
        cell = ws.cell(row=row_idx, column=col)
        cell.value = grand_total_count
        cell.alignment = center_alignment
        cell.font = Font(bold=True, size=12)
        cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")  # Light red
        
        # Grand Unique column
        cell = ws.cell(row=row_idx, column=col+1)
        cell.value = len(grand_unique_files)
        cell.alignment = center_alignment
        cell.font = Font(bold=True, size=12)
        cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")  # Light green
        
        # All Tools Used column
        all_tools = set()
        for run_data in data.values():
            all_tools.update(run_data["tools"])
        all_tools_str = ", ".join(sorted(all_tools))
        ws.cell(row=row_idx, column=col+2, value=all_tools_str)
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 12  # CWE ID
    ws.column_dimensions['B'].width = 60  # CWE Name
    
    # Each run has 2 columns (Total and Unique)
    col_idx = 3
    for _ in range(num_runs):
        col_letter_total = openpyxl.utils.get_column_letter(col_idx)
        col_letter_unique = openpyxl.utils.get_column_letter(col_idx + 1)
        ws.column_dimensions[col_letter_total].width = 12  # Total column
        ws.column_dimensions[col_letter_unique].width = 12  # Unique column
        col_idx += 2
    
    # Grand total columns and tools
    ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = 12  # Grand Total
    ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx + 1)].width = 12  # Grand Unique
    ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx + 2)].width = 25  # All Tools Used
    
    wb.save(out_path)


def scan_projects_background(files_info, clear_previous, scan_id):
    """Background task to scan projects and send progress updates"""
    global ALL_RUNS
    
    queue = progress_queues.get(scan_id)
    if not queue:
        return
    
    try:
        # Clear previous runs if requested
        if clear_previous:
            ALL_RUNS = []
            queue.put({"type": "info", "message": "Cleared previous runs"})
        
        total_files = len(files_info)
        
        # Process each uploaded file
        for idx, (filename, upload_path) in enumerate(files_info, 1):
            queue.put({
                "type": "progress",
                "current": idx,
                "total": total_files,
                "filename": filename,
                "stage": "extracting"
            })

            tmp_dir = Path(tempfile.mkdtemp(prefix="scanproj_"))

            try:
                # Extract
                queue.put({
                    "type": "status",
                    "message": f"[{idx}/{total_files}] Extracting {filename}..."
                })
                shutil.unpack_archive(str(upload_path), str(tmp_dir))
                
                # Delete the uploaded ZIP immediately after extraction
                try:
                    upload_path.unlink()
                    queue.put({
                        "type": "status",
                        "message": f"[{idx}/{total_files}] Cleaned up {filename}"
                    })
                except Exception as e:
                    print(f"Warning: Could not delete {upload_path}: {e}")

                # Scan
                queue.put({
                    "type": "progress",
                    "current": idx,
                    "total": total_files,
                    "filename": filename,
                    "stage": "scanning"
                })
                queue.put({
                    "type": "status",
                    "message": f"[{idx}/{total_files}] Scanning {filename} with 5 tools..."
                })
                
                # Run scanners (this is where most time is spent)
                def scanner_progress(message):
                    queue.put({
                        "type": "status",
                        "message": f"[{idx}/{total_files}] {message}"
                    })
                
                results = run_all_scanners(tmp_dir, progress_callback=scanner_progress)
                
                # Build summaries
                queue.put({
                    "type": "status",
                    "message": f"[{idx}/{total_files}] Building summaries for {filename}..."
                })
                sast_summary = build_sast_summary(results["sast"])
                dep_summary = build_dep_summary(results["dep"])

                project_name = filename.rsplit(".", 1)[0]

                # Store this run
                ALL_RUNS.append({
                    "scan_id": str(uuid.uuid4()),
                    "project_name": project_name,
                    "results": results,
                    "sast_summary": sast_summary,
                    "dep_summary": dep_summary,
                })
                
                queue.put({
                    "type": "status",
                    "message": f"[{idx}/{total_files}] ✅ Completed {filename}"
                })

            except Exception as e:
                queue.put({
                    "type": "error",
                    "message": f"[{idx}/{total_files}] ❌ Error scanning {filename}: {str(e)}"
                })

            finally:
                # Clean up temp extraction directory
                try:
                    shutil.rmtree(tmp_dir, ignore_errors=True)
                    if tmp_dir.exists():
                        # If rmtree failed, try force removal
                        import os
                        os.system(f"rm -rf {tmp_dir}")
                except Exception as e:
                    print(f"Warning: Could not remove temp dir {tmp_dir}: {e}")
        
        # Done
        queue.put({
            "type": "complete",
            "message": f"Successfully scanned {total_files} project(s)!",
            "total_scanned": total_files
        })
        
    except Exception as e:
        queue.put({
            "type": "error",
            "message": f"Fatal error: {str(e)}"
        })


def cleanup_old_reports():
    """Clean up report files older than 1 hour"""
    try:
        import time
        current_time = time.time()
        for file_path in REPORT_DIR.glob("*"):
            if file_path.is_file():
                # Get file age in seconds
                file_age = current_time - file_path.stat().st_mtime
                # Delete if older than 1 hour (3600 seconds)
                if file_age > 3600:
                    file_path.unlink()
                    print(f"Cleaned up old report: {file_path.name}")
    except Exception as e:
        print(f"Warning: Could not clean old reports: {e}")


def cleanup_temp_directories():
    """Clean up any leftover temp directories from crashed scans"""
    try:
        temp_base = Path(tempfile.gettempdir())
        for temp_dir in temp_base.glob("scanproj_*"):
            if temp_dir.is_dir():
                try:
                    shutil.rmtree(temp_dir)
                    print(f"Cleaned up temp directory: {temp_dir.name}")
                except:
                    pass
    except Exception as e:
        print(f"Warning: Could not clean temp directories: {e}")


@app.route("/", methods=["GET", "POST"])
def index():
    global ALL_RUNS

    if request.method == "POST":
        # Clean up old files before starting new scan
        cleanup_old_reports()
        cleanup_temp_directories()
        
        # Handle multiple file uploads
        files = request.files.getlist("project_zip")
        
        if not files or len(files) == 0:
            flash("No files uploaded", "warning")
            return redirect(request.url)
        
        # Filter valid files
        valid_files = [f for f in files if f.filename and allowed_file(f.filename)]
        
        if len(valid_files) == 0:
            flash("No valid .zip files uploaded", "warning")
            return redirect(request.url)
        
        # Save files and prepare info
        files_info = []
        for file in valid_files:
            filename = secure_filename(file.filename)
            upload_path = UPLOAD_DIR / filename
            file.save(upload_path)
            files_info.append((filename, upload_path))
        
        # Create scan ID and progress queue
        scan_id = str(uuid.uuid4())
        progress_queues[scan_id] = Queue()
        
        # Start background scanning
        clear_previous = request.form.get("clear_previous") == "yes"
        thread = Thread(target=scan_projects_background, args=(files_info, clear_previous, scan_id))
        thread.daemon = True
        thread.start()
        
        # Redirect to progress page
        return redirect(url_for("scan_progress", scan_id=scan_id))

    return render_template("index.html", num_runs=len(ALL_RUNS))


@app.route("/scan_progress/<scan_id>")
def scan_progress(scan_id):
    """Display progress page"""
    if scan_id not in progress_queues:
        flash("Invalid scan ID", "danger")
        return redirect(url_for("index"))
    
    return render_template("progress.html", scan_id=scan_id)


@app.route("/progress_stream/<scan_id>")
def progress_stream(scan_id):
    """Server-Sent Events stream for progress updates"""
    def generate():
        queue = progress_queues.get(scan_id)
        if not queue:
            yield f"data: {json.dumps({'type': 'error', 'message': 'Invalid scan ID'})}\n\n"
            return
        
        while True:
            try:
                # Get update from queue (timeout after 30 seconds)
                update = queue.get(timeout=30)
                yield f"data: {json.dumps(update)}\n\n"
                
                # If complete or error, clean up and stop
                if update.get("type") in ["complete", "error"]:
                    # Clean up queue after a delay
                    time.sleep(2)
                    if scan_id in progress_queues:
                        del progress_queues[scan_id]
                    break
                    
            except:
                # Timeout or error - send keepalive
                yield f"data: {json.dumps({'type': 'keepalive'})}\n\n"
    
    return Response(stream_with_context(generate()), mimetype="text/event-stream")


@app.route("/results")
def results_page():
    global ALL_RUNS
    if not ALL_RUNS:
        flash("No results yet. Upload projects first.", "warning")
        return redirect(url_for("index"))

    # Show the most recent run by default
    current_run = ALL_RUNS[-1]
    
    return render_template(
        "results.html",
        results=current_run["results"],
        sast_summary=current_run["sast_summary"],
        dep_summary=current_run["dep_summary"],
        project_name=current_run["project_name"],
        all_runs=ALL_RUNS,
        current_index=len(ALL_RUNS) - 1
    )


@app.route("/results/<int:run_index>")
def view_run(run_index):
    global ALL_RUNS
    
    if run_index < 0 or run_index >= len(ALL_RUNS):
        flash("Invalid run index", "danger")
        return redirect(url_for("results_page"))
    
    current_run = ALL_RUNS[run_index]
    
    return render_template(
        "results.html",
        results=current_run["results"],
        sast_summary=current_run["sast_summary"],
        dep_summary=current_run["dep_summary"],
        project_name=current_run["project_name"],
        all_runs=ALL_RUNS,
        current_index=run_index
    )


@app.route("/download_report/<int:run_index>")
def download_report(run_index):
    global ALL_RUNS
    
    if run_index < 0 or run_index >= len(ALL_RUNS):
        flash("Invalid run index", "danger")
        return redirect(url_for("results_page"))
    
    run_data = ALL_RUNS[run_index]
    project_name = run_data["project_name"]
    sast_summary = run_data["sast_summary"]
    dep_summary = run_data["dep_summary"]

    safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in project_name)
    filename = f"security_scan_report_{safe_name}.docx"
    out_path = REPORT_DIR / filename

    generate_docx_report(project_name, sast_summary, dep_summary, out_path)

    return send_from_directory(REPORT_DIR, filename, as_attachment=True)


@app.route("/download_comparison")
def download_comparison():
    """Download Excel with checkmarks (original format)"""
    global ALL_RUNS
    
    if len(ALL_RUNS) == 0:
        flash("No runs available to compare", "warning")
        return redirect(url_for("index"))
    
    filename = "cwe_comparison_checkmarks.xlsx"
    out_path = REPORT_DIR / filename
    
    generate_comparison_excel(ALL_RUNS, out_path)
    
    return send_from_directory(REPORT_DIR, filename, as_attachment=True)


@app.route("/download_detailed_counts")
def download_detailed_counts():
    """NEW: Download Excel with FILE COUNTS instead of checkmarks"""
    global ALL_RUNS
    
    if len(ALL_RUNS) == 0:
        flash("No runs available to compare", "warning")
        return redirect(url_for("index"))
    
    filename = "cwe_detailed_file_counts.xlsx"
    out_path = REPORT_DIR / filename
    
    generate_detailed_count_excel(ALL_RUNS, out_path)
    
    return send_from_directory(REPORT_DIR, filename, as_attachment=True)


@app.route("/download_json/<int:run_index>")
def download_json(run_index):
    """NEW: Download results as JSON in the user's specified format"""
    global ALL_RUNS
    
    if run_index < 0 or run_index >= len(ALL_RUNS):
        flash("Invalid run index", "danger")
        return redirect(url_for("results_page"))
    
    run_data = ALL_RUNS[run_index]
    project_name = run_data["project_name"]
    
    # Get RAW SAST results for accurate data
    sast_results = run_data.get("results", {}).get("sast", {})
    
    # Build JSON in user's format
    vulnerabilities = []
    vuln_id = 1
    
    for scanner_name, findings in sast_results.items():
        for finding in findings:
            cwe = finding.get("cwe", "CWE-UNKNOWN")
            if cwe == "CWE-UNKNOWN":
                continue
            
            vuln_entry = {
                "id": vuln_id,
                "cwe": cwe,
                "name": get_cwe_name(cwe),
                "file": finding.get("file", "unknown"),
                "line": finding.get("line", 0),
                "description": finding.get("message", f"Found by {scanner_name} scanner"),
                "severity": finding.get("severity", "UNKNOWN"),
                "scanner": scanner_name
            }
            vulnerabilities.append(vuln_entry)
            vuln_id += 1
    
    # Create final JSON structure
    output = {
        "run": project_name,
        "model": project_name,
        "language": "Multi-language",
        "vulnerabilities": vulnerabilities
    }
    
    # Save to file
    filename = f"vulnerabilities_{project_name}.json"
    out_path = REPORT_DIR / filename
    
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(output, f, indent=2)
    
    return send_from_directory(REPORT_DIR, filename, as_attachment=True)


def get_cwe_name(cwe: str) -> str:
    """Get human-readable CWE name - COMPREHENSIVE MAPPING (100+ CWEs)"""
    cwe_names = {
        # Input Validation
        "CWE-20": "Improper Input Validation",
        "CWE-74": "Improper Neutralization of Special Elements in Output Used by a Downstream Component ('Injection')",
        "CWE-75": "Failure to Sanitize Special Elements into a Different Plane (Special Element Injection)",
        "CWE-79": "Improper Neutralization of Input During Web Page Generation ('Cross-site Scripting')",
        "CWE-89": "Improper Neutralization of Special Elements used in an SQL Command ('SQL Injection')",
        "CWE-91": "XML Injection (aka Blind XPath Injection)",
        "CWE-94": "Improper Control of Generation of Code ('Code Injection')",
        "CWE-95": "Improper Neutralization of Directives in Dynamically Evaluated Code ('Eval Injection')",
        "CWE-96": "Improper Neutralization of Directives in Statically Saved Code ('Static Code Injection')",
        "CWE-97": "Improper Neutralization of Server-Side Includes (SSI) Within a Web Page",
        "CWE-98": "Improper Control of Filename for Include/Require Statement in PHP Program ('PHP Remote File Inclusion')",
        
        # Path Traversal & File Operations
        "CWE-22": "Improper Limitation of a Pathname to a Restricted Directory ('Path Traversal')",
        "CWE-23": "Relative Path Traversal",
        "CWE-36": "Absolute Path Traversal",
        "CWE-73": "External Control of File Name or Path",
        "CWE-434": "Unrestricted Upload of File with Dangerous Type",
        "CWE-59": "Improper Link Resolution Before File Access ('Link Following')",
        "CWE-377": "Insecure Temporary File",
        
        # Command Injection
        "CWE-78": "Improper Neutralization of Special Elements used in an OS Command ('OS Command Injection')",
        "CWE-77": "Improper Neutralization of Special Elements used in a Command ('Command Injection')",
        
        # Authentication & Access Control
        "CWE-287": "Improper Authentication",
        "CWE-306": "Missing Authentication for Critical Function",
        "CWE-285": "Improper Authorization",
        "CWE-284": "Improper Access Control",
        "CWE-862": "Missing Authorization",
        "CWE-863": "Incorrect Authorization",
        "CWE-276": "Incorrect Default Permissions",
        "CWE-732": "Incorrect Permission Assignment for Critical Resource",
        "CWE-250": "Execution with Unnecessary Privileges",
        "CWE-269": "Improper Privilege Management",
        
        # Cryptography
        "CWE-327": "Use of a Broken or Risky Cryptographic Algorithm",
        "CWE-328": "Use of Weak Hash",
        "CWE-326": "Inadequate Encryption Strength",
        "CWE-321": "Use of Hard-coded Cryptographic Key",
        "CWE-322": "Key Exchange without Entity Authentication",
        "CWE-323": "Reusing a Nonce, Key Pair in Encryption",
        "CWE-324": "Use of a Key Past its Expiration Date",
        "CWE-325": "Missing Cryptographic Step",
        "CWE-329": "Generation of Predictable IV with CBC Mode",
        "CWE-338": "Use of Cryptographically Weak Pseudo-Random Number Generator (PRNG)",
        "CWE-330": "Use of Insufficiently Random Values",
        "CWE-331": "Insufficient Entropy",
        "CWE-335": "Incorrect Usage of Seeds in Pseudo-Random Number Generator (PRNG)",
        "CWE-336": "Same Seed in Pseudo-Random Number Generator (PRNG)",
        "CWE-337": "Predictable Seed in Pseudo-Random Number Generator (PRNG)",
        
        # SSL/TLS
        "CWE-295": "Improper Certificate Validation",
        "CWE-296": "Improper Following of a Certificate's Chain of Trust",
        "CWE-297": "Improper Validation of Certificate with Host Mismatch",
        "CWE-298": "Improper Validation of Certificate Expiration",
        "CWE-299": "Improper Check for Certificate Revocation",
        
        # Sensitive Data
        "CWE-200": "Exposure of Sensitive Information to an Unauthorized Actor",
        "CWE-209": "Generation of Error Message Containing Sensitive Information",
        "CWE-215": "Insertion of Sensitive Information Into Debugging Code",
        "CWE-312": "Cleartext Storage of Sensitive Information",
        "CWE-313": "Cleartext Storage in a File or on Disk",
        "CWE-314": "Cleartext Storage in the Registry",
        "CWE-315": "Cleartext Storage of Sensitive Information in a Cookie",
        "CWE-316": "Cleartext Storage of Sensitive Information in Memory",
        "CWE-317": "Cleartext Storage of Sensitive Information in GUI",
        "CWE-318": "Cleartext Storage of Sensitive Information in Executable",
        "CWE-319": "Cleartext Transmission of Sensitive Information",
        "CWE-798": "Use of Hard-coded Credentials",
        "CWE-259": "Use of Hard-coded Password",
        "CWE-257": "Storing Passwords in a Recoverable Format",
        
        # Session Management
        "CWE-384": "Session Fixation",
        "CWE-613": "Insufficient Session Expiration",
        "CWE-614": "Sensitive Cookie in HTTPS Session Without 'Secure' Attribute",
        "CWE-1004": "Sensitive Cookie Without 'HttpOnly' Flag",
        "CWE-565": "Reliance on Cookies without Validation and Integrity Checking",
        
        # CSRF & Redirects
        "CWE-352": "Cross-Site Request Forgery (CSRF)",
        "CWE-601": "URL Redirection to Untrusted Site ('Open Redirect')",
        
        # XML & XXE
        "CWE-611": "Improper Restriction of XML External Entity Reference",
        "CWE-827": "Improper Control of Document Type Definition",
        
        # Deserialization
        "CWE-502": "Deserialization of Untrusted Data",
        "CWE-915": "Improperly Controlled Modification of Dynamically-Determined Object Attributes",
        
        # Resource Management
        "CWE-400": "Uncontrolled Resource Consumption",
        "CWE-404": "Improper Resource Shutdown or Release",
        "CWE-770": "Allocation of Resources Without Limits or Throttling",
        "CWE-771": "Missing Reference to Active Allocated Resource",
        "CWE-772": "Missing Release of Resource after Effective Lifetime",
        "CWE-775": "Missing Release of File Descriptor or Handle after Effective Lifetime",
        "CWE-776": "Unrestricted Recursion",
        "CWE-834": "Excessive Iteration",
        
        # Memory Safety
        "CWE-119": "Improper Restriction of Operations within the Bounds of a Memory Buffer",
        "CWE-120": "Buffer Copy without Checking Size of Input ('Classic Buffer Overflow')",
        "CWE-121": "Stack-based Buffer Overflow",
        "CWE-122": "Heap-based Buffer Overflow",
        "CWE-125": "Out-of-bounds Read",
        "CWE-787": "Out-of-bounds Write",
        "CWE-416": "Use After Free",
        "CWE-415": "Double Free",
        "CWE-476": "NULL Pointer Dereference",
        "CWE-401": "Missing Release of Memory after Effective Lifetime",
        "CWE-911": "Improper Update of Reference Count",
        
        # Integer & Numeric Errors
        "CWE-190": "Integer Overflow or Wraparound",
        "CWE-191": "Integer Underflow (Wrap or Wraparound)",
        "CWE-680": "Integer Overflow to Buffer Overflow",
        "CWE-681": "Incorrect Conversion between Numeric Types",
        "CWE-682": "Incorrect Calculation",
        "CWE-369": "Divide By Zero",
        
        # Race Conditions
        "CWE-362": "Concurrent Execution using Shared Resource with Improper Synchronization ('Race Condition')",
        "CWE-367": "Time-of-check Time-of-use (TOCTOU) Race Condition",
        "CWE-364": "Signal Handler Race Condition",
        
        # Logic Errors
        "CWE-670": "Always-Incorrect Control Flow Implementation",
        "CWE-571": "Expression is Always True",
        "CWE-570": "Expression is Always False",
        "CWE-561": "Dead Code",
        "CWE-489": "Active Debug Code",
        "CWE-501": "Trust Boundary Violation",
        
        # Logging & Monitoring
        "CWE-117": "Improper Output Neutralization for Logs",
        "CWE-532": "Insertion of Sensitive Information into Log File",
        "CWE-533": "DEPRECATED: Information Exposure Through Server Log Files",
        
        # Configuration
        "CWE-1188": "Insecure Default Initialization of Resource",
        "CWE-426": "Untrusted Search Path",
        "CWE-427": "Uncontrolled Search Path Element",
        "CWE-829": "Inclusion of Functionality from Untrusted Control Sphere",
        "CWE-830": "Inclusion of Web Functionality from an Untrusted Source",
        
        # Code Quality
        "CWE-477": "Use of Obsolete Function",
        "CWE-478": "Missing Default Case in Multiple Condition Expression",
        "CWE-479": "Signal Handler Use of a Non-reentrant Function",
        "CWE-480": "Use of Incorrect Operator",
        "CWE-483": "Incorrect Block Delimitation",
        "CWE-484": "Omitted Break Statement in Switch",
        
        # Expression Language Injection
        "CWE-917": "Improper Neutralization of Special Elements used in an Expression Language Statement ('Expression Language Injection')",
        
        # Password & Auth
        "CWE-521": "Weak Password Requirements",
        "CWE-916": "Use of Password Hash With Insufficient Computational Effort",
        
        # Regex
        "CWE-1333": "Inefficient Regular Expression Complexity",
        
        # Prototype Pollution
        "CWE-1321": "Improperly Controlled Modification of Object Prototype Attributes ('Prototype Pollution')",
        
        # Server-Side Request Forgery
        "CWE-918": "Server-Side Request Forgery (SSRF)",
        
        # Null Byte Injection
        "CWE-158": "Improper Neutralization of Null Byte or NUL Character",
        
        # Format String
        "CWE-134": "Use of Externally-Controlled Format String",
        
        # Information Disclosure
        "CWE-203": "Observable Discrepancy",
        "CWE-208": "Observable Timing Discrepancy",
        
        # Missing Support
        "CWE-353": "Missing Support for Integrity Check",
        
        # ZIP vulnerabilities
        "CWE-409": "Improper Handling of Highly Compressed Data (Data Amplification)",
        "CWE-410": "Insufficient Resource Pool",
    }
    return cwe_names.get(cwe, f"Unknown Vulnerability ({cwe})")


@app.route("/clear_runs")
def clear_runs():
    global ALL_RUNS
    ALL_RUNS = []
    flash("All previous runs cleared", "info")
    return redirect(url_for("index"))


@app.route("/cleanup_all")
def cleanup_all():
    """Manual cleanup endpoint - removes all temporary files"""
    try:
        # Clean reports
        cleanup_old_reports()
        
        # Clean temp directories
        cleanup_temp_directories()
        
        # Clean uploads
        for file in UPLOAD_DIR.glob("*.zip"):
            file.unlink()
        
        # Force clean ALL reports (not just old ones)
        for file in REPORT_DIR.glob("*"):
            if file.is_file():
                file.unlink()
        
        flash("✅ All temporary files cleaned up!", "success")
    except Exception as e:
        flash(f"Warning: Some files could not be cleaned: {e}", "warning")
    
    return redirect(url_for("index"))


if __name__ == "__main__":
    print("=" * 60)
    print("Multi-Run Security Scanner Starting...")
    print("=" * 60)
    
    # Cleanup on startup
    print("\n🧹 Cleaning up old files...")
    cleanup_old_reports()
    cleanup_temp_directories()
    
    # Clean uploads directory
    try:
        for file in UPLOAD_DIR.glob("*.zip"):
            file.unlink()
            print(f"Removed old upload: {file.name}")
    except Exception as e:
        print(f"Warning: Could not clean uploads: {e}")
    
    print("\n✅ Cleanup complete!")
    print("\n🚀 Starting server on http://localhost:8080")
    print("=" * 60)
    print()
    
    app.run(host="0.0.0.0", port=8080, debug=True, use_reloader=False)
